"""
Document parsing utilities.
Extracts raw text from PDF and DOCX files, then cleans/normalises it.
"""

import re
import io
import logging
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


# ── Public API ──────────────────────────────────────────────────────────────

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns cleaned, normalised plain text.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        raw = _extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        raw = _extract_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return clean_text(raw)


def clean_text(text: str) -> str:
    """Remove noise and normalise whitespace."""
    # Remove non-printable characters (keep newlines/tabs)
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
    # Collapse multiple spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse more than 2 consecutive newlines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Parsers ─────────────────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber (handles tables & columns)."""
    pages: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if page_text:
                    pages.append(page_text)
    except Exception as exc:
        logger.warning("pdfplumber failed (%s), text may be incomplete.", exc)
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX, including tables."""
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ── Section splitter (best-effort) ─────────────────────────────────────────

SECTION_HEADERS = {
    "skills": r"(?i)(technical\s+skills?|skills?|core\s+competencies|technologies)",
    "experience": r"(?i)(work\s+experience|experience|employment|professional\s+background)",
    "education": r"(?i)(education|academic|qualifications?|degrees?)",
    "projects": r"(?i)(projects?|personal\s+projects?|key\s+projects?|portfolio)",
    "certifications": r"(?i)(certifications?|licen[sc]es?|achievements?)",
    "summary": r"(?i)(summary|objective|profile|about\s+me)",
}


def split_into_sections(text: str) -> dict[str, str]:
    """
    Attempt to split resume text into labelled sections.
    Falls back to returning the whole text under 'full_text'.
    """
    lines = text.split("\n")
    sections: dict[str, list[str]] = {"full_text": []}
    current_section = "full_text"

    for line in lines:
        matched = False
        for section_name, pattern in SECTION_HEADERS.items():
            # A line is a header if it matches the pattern and is short (<= 60 chars)
            if re.search(pattern, line) and len(line.strip()) <= 60:
                current_section = section_name
                matched = True
                break
        if not matched:
            sections.setdefault(current_section, []).append(line)

    return {k: "\n".join(v).strip() for k, v in sections.items() if v}
